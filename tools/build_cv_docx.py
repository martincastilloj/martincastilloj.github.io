from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
from reportlab.platypus.flowables import HRFlowable


ROOT = Path(__file__).resolve().parents[1]
DOCX_OUT = ROOT / "assets" / "docs" / "cv-martin-castillo.docx"
PDF_OUT = ROOT / "assets" / "docs" / "cv-martin-castillo.pdf"


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(28, 33, 38)
MUTED = RGBColor(90, 96, 105)


PROFILE = (
    "Estudiante de Ingeniería Comercial y Magíster en Economía y Políticas Públicas "
    "en la Universidad Adolfo Ibáñez, con experiencia en investigación aplicada, "
    "docencia universitaria e iniciativas estudiantiles orientadas al análisis público. "
    "Sus intereses se concentran en economía política, políticas públicas comparadas, "
    "gobernanza, relaciones internacionales e historia económica, con especial atención "
    "al Asia-Pacífico."
)

EDUCATION = [
    ("Universidad Adolfo Ibáñez", "Ingeniería Comercial, 2021-2026."),
    ("Universidad Adolfo Ibáñez", "Magíster en Economía y Políticas Públicas, 2025-2026."),
    ("Universidad Adolfo Ibáñez", "Licenciado en Ciencias Económicas, marzo 2021-diciembre 2024."),
]

EXPERIENCE = [
    (
        "Presidente y co-fundador",
        "CEAPS | Club de Economía, Actualidad Política y Sociedad",
        "mayo 2026-actualidad",
        [
            "Coordinación estratégica de una organización universitaria dedicada al análisis de economía, política, sociedad y actualidad nacional e internacional.",
            "Organización de charlas académicas, talleres formativos, visitas institucionales, publicaciones estudiantiles y espacios de discusión.",
            "Vinculación con académicos, estudiantes, centros de estudio e instituciones públicas.",
        ],
    ),
    (
        "Práctica profesional como ayudante de investigación",
        "GobLab UAI",
        "abril 2025-julio 2025",
        [
            "Investigación y selección de algoritmos implementados por el sector público en Chile para su incorporación en el Repositorio de Algoritmos Públicos.",
            "Elaboración de fichas técnicas, apoyo en análisis de información, redacción de informes y logística de eventos.",
        ],
    ),
    (
        "Ayudante de investigación",
        "Escuela de Gobierno UAI",
        "octubre 2025-noviembre 2025",
        [
            "Recolección y sistematización de datos sobre cumplimiento de la ley de transparencia en municipalidades chilenas.",
            "Construcción de bases de datos a partir de información publicada por municipios y registros del Consejo para la Transparencia.",
        ],
    ),
]

TEACHING_BULLETS = [
    "Ayudantías de cátedra, laboratorio y corrección en economía, métodos cuantitativos, ciencia política, relaciones internacionales e historia de Asia.",
    "Cursos: Microeconometría, Econometría I, Descripción y Visualización de Datos, Razonamiento Cuantitativo con Datos I, Contabilidad y Crisis de las Democracias.",
    "Cursos Asia-Pacífico: Historia Política de China, Historia de las relaciones de Chile con China, Relaciones Internacionales de Chile con China y el Asia Pacífico, Historia del Asia Oriental e Historia y Sociedad de Japón.",
]

COMPACT_LINES = [
    ("Columna", '"¿Está muriendo la Universidad?", publicada originalmente en CEAPS, 23 de mayo de 2026.'),
    ("Herramientas", "Microsoft Excel, Stata, RStudio y análisis de información pública institucional."),
    ("Intereses", "Economía política, políticas públicas comparadas, relaciones internacionales, historia económica, Asia-Pacífico, China, Japón, gobernanza e instituciones."),
    ("Idiomas", "Español nativo; inglés avanzado."),
]


def set_run(run, *, bold=False, italic=False, size=10.5, color=INK):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color


def set_paragraph(p, *, before=0, after=4, line=1.15):
    fmt = p.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_heading(doc, text):
    p = doc.add_paragraph()
    set_paragraph(p, before=10, after=4, line=1.0)
    r = p.add_run(text.upper())
    set_run(r, bold=True, size=11.5, color=BLUE)
    return p


def add_role(doc, title, org, dates, bullets):
    p = doc.add_paragraph()
    set_paragraph(p, before=3, after=1, line=1.05)
    r = p.add_run(title)
    set_run(r, bold=True, size=10.5, color=INK)
    r = p.add_run(f" | {org}")
    set_run(r, size=10.5, color=INK)
    r = p.add_run(f" | {dates}")
    set_run(r, italic=True, size=9.8, color=MUTED)

    for bullet in bullets:
        bp = doc.add_paragraph(style="List Bullet")
        set_paragraph(bp, before=0, after=2, line=1.12)
        set_run(bp.add_run(bullet), size=9.8, color=INK)


def add_compact_line(doc, label, text):
    p = doc.add_paragraph()
    set_paragraph(p, before=0, after=2, line=1.1)
    set_run(p.add_run(label + ": "), bold=True, size=9.8, color=DARK_BLUE)
    set_run(p.add_run(text), size=9.8, color=INK)


def add_bottom_border(paragraph, color="D9DEE6"):
    p = paragraph._p
    p_pr = p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def main():
    build_docx()
    build_pdf()
    print(DOCX_OUT)
    print(PDF_OUT)


def build_docx():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"].font.color.rgb = INK
    styles["List Bullet"].font.name = "Calibri"
    styles["List Bullet"].font.size = Pt(9.8)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph(title, before=0, after=1, line=1.0)
    set_run(title.add_run("MARTÍN CASTILLO JIMÉNEZ"), bold=True, size=18, color=DARK_BLUE)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph(subtitle, before=0, after=2, line=1.05)
    set_run(
        subtitle.add_run(
            "Economía, políticas públicas, relaciones internacionales e investigación aplicada"
        ),
        size=10.5,
        color=MUTED,
    )

    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph(contact, before=0, after=7, line=1.05)
    set_run(
        contact.add_run(
            "Santiago de Chile | martin.guillermo.j@gmail.com | linkedin.com/in/martincastilloj | martincastilloj.github.io"
        ),
        size=9.2,
        color=INK,
    )
    add_bottom_border(contact)

    add_heading(doc, "Perfil")
    p = doc.add_paragraph()
    set_paragraph(p, before=0, after=4, line=1.18)
    set_run(
        p.add_run(
            PROFILE
        ),
        size=10.2,
        color=INK,
    )

    add_heading(doc, "Formación")
    for label, text in EDUCATION:
        add_compact_line(doc, label, text)

    add_heading(doc, "Experiencia")
    for title, org, dates, bullets in EXPERIENCE:
        add_role(doc, title, org, dates, bullets)

    add_heading(doc, "Docencia universitaria")
    add_role(
        doc,
        "Ayudante universitario",
        "Universidad Adolfo Ibáñez",
        "2023-actualidad",
        TEACHING_BULLETS,
    )

    add_heading(doc, "Escritos")
    add_compact_line(doc, *COMPACT_LINES[0])

    add_heading(doc, "Habilidades e intereses")
    for label, text in COMPACT_LINES[1:]:
        add_compact_line(doc, label, text)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(footer.add_run("CV generado para martincastilloj.github.io | Mayo 2026"), size=8.5, color=MUTED)

    DOCX_OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(DOCX_OUT)


def build_pdf():
    styles = getSampleStyleSheet()
    doc = SimpleDocTemplate(
        str(PDF_OUT),
        pagesize=LETTER,
        rightMargin=1.65 * cm,
        leftMargin=1.65 * cm,
        topMargin=1.45 * cm,
        bottomMargin=1.45 * cm,
        title="CV - Martin Castillo Jimenez",
        author="Martin Castillo Jimenez",
    )

    title_style = ParagraphStyle(
        "Title",
        parent=styles["Title"],
        fontName="Helvetica-Bold",
        fontSize=18,
        leading=21,
        textColor=colors.HexColor("#1F4D78"),
        alignment=1,
        spaceAfter=2,
    )
    subtitle_style = ParagraphStyle(
        "Subtitle",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=9.7,
        leading=12,
        textColor=colors.HexColor("#5A6069"),
        alignment=1,
        spaceAfter=3,
    )
    section_style = ParagraphStyle(
        "Section",
        parent=styles["Heading2"],
        fontName="Helvetica-Bold",
        fontSize=10.5,
        leading=12,
        textColor=colors.HexColor("#2E74B5"),
        spaceBefore=9,
        spaceAfter=4,
        uppercase=True,
    )
    body_style = ParagraphStyle(
        "Body",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=9.3,
        leading=11.5,
        textColor=colors.HexColor("#1C2126"),
        spaceAfter=4,
    )
    role_style = ParagraphStyle(
        "Role",
        parent=body_style,
        fontName="Helvetica",
        fontSize=9.4,
        leading=11.5,
        spaceBefore=3,
        spaceAfter=1,
    )
    bullet_style = ParagraphStyle(
        "Bullet",
        parent=body_style,
        leftIndent=12,
        firstLineIndent=-7,
        spaceAfter=2,
    )

    story = [
        Paragraph("MARTÍN CASTILLO JIMÉNEZ", title_style),
        Paragraph(
            "Economía, políticas públicas, relaciones internacionales e investigación aplicada",
            subtitle_style,
        ),
        Paragraph(
            "Santiago de Chile | martin.guillermo.j@gmail.com | linkedin.com/in/martincastilloj | martincastilloj.github.io",
            subtitle_style,
        ),
        HRFlowable(width="100%", thickness=0.6, color=colors.HexColor("#D9DEE6"), spaceBefore=2, spaceAfter=8),
        Paragraph("PERFIL", section_style),
        Paragraph(PROFILE, body_style),
        Paragraph("FORMACIÓN", section_style),
    ]

    for label, text in EDUCATION:
        story.append(Paragraph(f"<b>{label}:</b> {text}", body_style))

    story.append(Paragraph("EXPERIENCIA", section_style))
    for title, org, dates, bullets in EXPERIENCE:
        story.append(Paragraph(f"<b>{title}</b> | {org} | <i>{dates}</i>", role_style))
        for bullet in bullets:
            story.append(Paragraph(f"• {bullet}", bullet_style))

    story.append(Paragraph("DOCENCIA UNIVERSITARIA", section_style))
    story.append(Paragraph("<b>Ayudante universitario</b> | Universidad Adolfo Ibáñez | <i>2023-actualidad</i>", role_style))
    for bullet in TEACHING_BULLETS:
        story.append(Paragraph(f"• {bullet}", bullet_style))

    story.append(Paragraph("ESCRITOS", section_style))
    story.append(Paragraph(f"<b>{COMPACT_LINES[0][0]}:</b> {COMPACT_LINES[0][1]}", body_style))

    story.append(Paragraph("HABILIDADES E INTERESES", section_style))
    for label, text in COMPACT_LINES[1:]:
        story.append(Paragraph(f"<b>{label}:</b> {text}", body_style))

    story.append(Spacer(1, 4))
    story.append(Paragraph("CV generado para martincastilloj.github.io | Mayo 2026", subtitle_style))
    PDF_OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.build(story)


if __name__ == "__main__":
    main()
